"""``create_report`` — compose an A4 report (titled sections, paragraph
explanations, optional charts) and save it as both a .docx and a .pdf into
the session workspace.

Backs the desktop "Report Generator" tools-rail entry. Charts, when present,
are rendered once via matplotlib and embedded into both output formats, so
the .docx and .pdf are two independent renderers reading one shared content
model rather than a docx->pdf conversion (avoids a LibreOffice/Java-class
dependency, consistent with tools/read_extract.py's opendataloader-pdf
decision this session).

All three libraries (python-docx, matplotlib, reportlab) live in the
``documents`` extra (see pyproject.toml) and import lazily here, so a
partial install degrades to a clear tool_error instead of crashing the
agent loop.
"""

from __future__ import annotations

import json
import re
import shutil
import tempfile
from pathlib import Path
from typing import Any

from tools.file_tools import _check_workspace_selected, _resolve_path_for_task
from tools.registry import registry, tool_error, tool_result

__all__ = ["create_report_tool"]

_A4_WIDTH_MM = 210
_A4_HEIGHT_MM = 297
_CHART_TYPES = frozenset({"bar", "line", "pie"})
_MAX_SECTIONS = 60


def _sanitize_filename_base(name: str) -> str:
    """Same hygiene as _sanitize_attachment_name in tui_gateway/server.py:
    basename-only (no path traversal), strip control chars, never empty."""
    candidate = Path(str(name or "").strip()).name
    candidate = re.sub(r"[\x00-\x1f]+", "_", candidate)
    candidate = candidate.strip().strip(".")
    return candidate or "report"


def _require(module_name: str, pip_name: str):
    try:
        return __import__(module_name)
    except ImportError as exc:
        raise RuntimeError(
            f"{pip_name} is not installed. Re-run the Pocura installer "
            f"(scripts/install.ps1 / install.sh) to get the document-generation "
            f"libraries, or install `{pip_name}` manually."
        ) from exc


def _validate_sections(sections: Any) -> str | None:
    if not isinstance(sections, list) or not sections:
        return "create_report: 'sections' must be a non-empty array."
    if len(sections) > _MAX_SECTIONS:
        return f"create_report: too many sections ({len(sections)}); max is {_MAX_SECTIONS}."
    for i, section in enumerate(sections):
        if not isinstance(section, dict):
            return f"create_report: sections[{i}] must be an object."
        if not str(section.get("heading") or "").strip():
            return f"create_report: sections[{i}] is missing a non-empty 'heading'."
        if not isinstance(section.get("body"), str) or not section["body"].strip():
            return f"create_report: sections[{i}] is missing non-empty 'body' text."
        chart = section.get("chart")
        if chart is None:
            continue
        if not isinstance(chart, dict):
            return f"create_report: sections[{i}].chart must be an object."
        if chart.get("type") not in _CHART_TYPES:
            return f"create_report: sections[{i}].chart.type must be one of {sorted(_CHART_TYPES)}."
        labels, values = chart.get("labels"), chart.get("values")
        if not isinstance(labels, list) or not isinstance(values, list) or not labels or not values:
            return f"create_report: sections[{i}].chart needs non-empty 'labels' and 'values' arrays."
        if len(labels) != len(values):
            return f"create_report: sections[{i}].chart 'labels' and 'values' must be the same length."
        if not all(isinstance(v, (int, float)) for v in values):
            return f"create_report: sections[{i}].chart 'values' must all be numbers."
    return None


def _render_chart(chart: dict, out_path: Path, matplotlib_mod) -> None:
    # Agg is the headless, GUI-less backend -- must be selected before pyplot
    # is imported anywhere in the process, hence the local import here rather
    # than at module load (a different tool importing pyplot first would
    # otherwise lock in whatever default backend matplotlib picked).
    matplotlib_mod.use("Agg", force=False)
    import matplotlib.pyplot as plt

    labels = [str(label) for label in chart["labels"]]
    values = [float(v) for v in chart["values"]]
    chart_title = str(chart.get("chart_title") or "").strip()

    fig, ax = plt.subplots(figsize=(6.5, 4), dpi=150)
    try:
        kind = chart["type"]
        if kind == "bar":
            ax.bar(labels, values)
            ax.tick_params(axis="x", rotation=30)
        elif kind == "line":
            ax.plot(labels, values, marker="o")
            ax.tick_params(axis="x", rotation=30)
        else:  # pie
            ax.pie(values, labels=labels, autopct="%1.1f%%")
            ax.axis("equal")
        if chart_title:
            ax.set_title(chart_title)
        fig.tight_layout()
        fig.savefig(str(out_path))
    finally:
        plt.close(fig)


def _build_docx(title: str, sections: list[dict], chart_paths: dict[int, Path], out_path: Path) -> None:
    docx_mod = _require("docx", "python-docx")
    from docx.shared import Mm

    doc = docx_mod.Document()
    section = doc.sections[0]
    section.page_width = Mm(_A4_WIDTH_MM)
    section.page_height = Mm(_A4_HEIGHT_MM)

    doc.add_heading(title, level=0)
    for i, block in enumerate(sections):
        doc.add_heading(str(block["heading"]), level=1)
        doc.add_paragraph(str(block["body"]))
        if i in chart_paths:
            doc.add_picture(str(chart_paths[i]), width=Mm(_A4_WIDTH_MM - 60))

    doc.save(str(out_path))


def _build_pdf(title: str, sections: list[dict], chart_paths: dict[int, Path], out_path: Path) -> None:
    # _require()'s __import__(name) returns the TOP-LEVEL package for a dotted
    # name (reportlab, not reportlab.lib.pagesizes) -- fine for the single-name
    # docx/matplotlib checks above, wrong here, so import directly instead.
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Image as RLImage
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError(
            "reportlab is not installed. Re-run the Pocura installer "
            "(scripts/install.ps1 / install.sh) to get the document-generation "
            "libraries, or install `reportlab` manually."
        ) from exc
    from xml.sax.saxutils import escape

    styles = getSampleStyleSheet()
    story = [Paragraph(escape(title), styles["Title"]), Spacer(1, 14)]

    for i, block in enumerate(sections):
        story.append(Paragraph(escape(str(block["heading"])), styles["Heading1"]))
        # reportlab's Paragraph takes a tiny XML-like markup, not plain text --
        # unescaped '&'/'<'/'>' in the body would break its parser outright.
        for line in str(block["body"]).split("\n"):
            if line.strip():
                story.append(Paragraph(escape(line), styles["BodyText"]))
        if i in chart_paths:
            story.append(Spacer(1, 8))
            story.append(RLImage(str(chart_paths[i]), width=400, height=246))
        story.append(Spacer(1, 14))

    SimpleDocTemplate(str(out_path), pagesize=A4).build(story)


def create_report_tool(
    title: str,
    sections: list[dict],
    filename_base: str,
    task_id: str = "default",
) -> str:
    workspace_err = _check_workspace_selected(task_id)
    if workspace_err:
        return tool_error(workspace_err)

    if not isinstance(title, str) or not title.strip():
        return tool_error("create_report: 'title' must be a non-empty string.")
    validation_err = _validate_sections(sections)
    if validation_err:
        return tool_error(validation_err)

    safe_name = _sanitize_filename_base(filename_base)
    try:
        docx_path = _resolve_path_for_task(f"{safe_name}.docx", task_id)
        pdf_path = _resolve_path_for_task(f"{safe_name}.pdf", task_id)
    except Exception as exc:  # noqa: BLE001 - surfaced as a plain tool error below
        return tool_error(f"create_report: could not resolve output path: {exc}")

    tmp_dir = Path(tempfile.mkdtemp(prefix="pocura-report-"))
    try:
        chart_paths: dict[int, Path] = {}
        try:
            matplotlib_mod = _require("matplotlib", "matplotlib")
            for i, block in enumerate(sections):
                chart = block.get("chart")
                if chart:
                    chart_png = tmp_dir / f"chart_{i}.png"
                    _render_chart(chart, chart_png, matplotlib_mod)
                    chart_paths[i] = chart_png
        except RuntimeError as exc:
            if any(block.get("chart") for block in sections):
                return tool_error(str(exc))
            # No section actually asked for a chart -- a missing chart library
            # is irrelevant, don't block a report that has none.

        try:
            _build_docx(title, sections, chart_paths, docx_path)
            _build_pdf(title, sections, chart_paths, pdf_path)
        except RuntimeError as exc:
            return tool_error(str(exc))
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    return tool_result(
        docx_path=str(docx_path),
        pdf_path=str(pdf_path),
        title=title,
        section_count=len(sections),
    )


def check_report_requirements() -> bool:
    """Always listed -- a missing optional library degrades to a clear
    tool_error at call time (matches tools/read_extract.py's contract),
    rather than silently hiding the tool from the model."""
    return True


CREATE_REPORT_SCHEMA = {
    "name": "create_report",
    "description": (
        "Generate an A4-paginated report (title + titled sections, each with a "
        "paragraph explanation and an optional chart) and save it into the "
        "workspace as BOTH a .docx and a .pdf. Call this once, after reading "
        "every attached source and drafting the full report content -- it "
        "produces both file formats together in one call.\n\n"
        "Only add a chart to a section when the sources contain real numeric "
        "data suited to it; never invent numbers to populate one."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "title": {
                "type": "string",
                "description": "Report title.",
            },
            "sections": {
                "type": "array",
                "description": "Report sections, in order.",
                "items": {
                    "type": "object",
                    "properties": {
                        "heading": {"type": "string", "description": "Section heading."},
                        "body": {
                            "type": "string",
                            "description": "Paragraph explanation for this section (plain text).",
                        },
                        "chart": {
                            "type": "object",
                            "description": "Optional chart built from real numeric data found in the sources.",
                            "properties": {
                                "type": {
                                    "type": "string",
                                    "enum": ["bar", "line", "pie"],
                                },
                                "labels": {
                                    "type": "array",
                                    "items": {"type": "string"},
                                },
                                "values": {
                                    "type": "array",
                                    "items": {"type": "number"},
                                },
                                "chart_title": {"type": "string"},
                            },
                            "required": ["type", "labels", "values"],
                        },
                    },
                    "required": ["heading", "body"],
                },
            },
            "filename_base": {
                "type": "string",
                "description": "Base filename (without extension) for the output files.",
            },
        },
        "required": ["title", "sections", "filename_base"],
    },
}


registry.register(
    name="create_report",
    toolset="file",
    schema=CREATE_REPORT_SCHEMA,
    handler=lambda args, **kw: create_report_tool(
        title=args.get("title", ""),
        sections=args.get("sections") or [],
        filename_base=args.get("filename_base", ""),
        task_id=kw.get("task_id") or "default",
    ),
    check_fn=check_report_requirements,
    emoji="📊",
)
